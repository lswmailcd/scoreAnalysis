# -*- coding: utf-8 -*-

import numpy as np
import xlrd
import matplotlib.pyplot as plt
from pylab import mpl
from docx import Document
import sys
sys.path.append('../stockAnalysis')
import Graph as gp
#======================================================================================
#需要根据具体分析报告修改的数据
#scoreAnalysisTemplate中各小班编号从0开始，依次为0|1|2|3|4|...，不能出现中文字
excelSourceDataFileName = u'.\\data\\2019级眼视光医学10合班医用高等数学成绩.xls' #scoreTemplate.xls'
stuNum = 92
#以下数据起始标号为0
scoreColumnIdx = 3
classNameColumnIdx = 2
startRowIdx = 0
wordGraph_Title=u"2019级眼视光医学10合班“医用高等数学”成绩直方图"
yAxiScaleStep = 5 #每5分一个Y刻度
#======================================================================================
wordTemplateFileName = u'.\\data\\scoreAnalysisTemplate_math.docx'
wordGraph_XLable = u'成绩段'
wordGraph_YLable = u'学生人数'
scorePhasePlot = [30,35,40,45,50,55,60,65,70,75,80,85,90,95,100]
xtickLable = ['<30', '<35', '<40', '<45', '<50', '<55', '<60',
              '<65', '<70', '<75', '<80', '<85', '<90', '<95','<=100']
scorePhaseString = [u'30以下', u'30~34', u'35~39', u'40~44', u'45~49', u'50~54',u'55~59',
                    u'60~64',u'65~69', u'70~74', u'75~79', u'80~84', u'85~89', u'90~94', u'95~100']
wordGraph_FontSize = 25
#====================================================================================
data = xlrd.open_workbook(excelSourceDataFileName)
tableXLS = data.sheets()[0]
nStuTakExam = tableXLS.nrows-startRowIdx
nDelayExamStu = stuNum-nStuTakExam
nNotAttendExamStu= 0 #旷考人数
nViolationExamStu= 0 #违纪人数
nCheatExamStu= 0 #作弊人数
print "学生总人数为:%r,  参考学生人数为：%r,  缓考学生人数为：%r" % (stuNum, nStuTakExam, nDelayExamStu)
nMaxClassNum = nStuTakExam/20
score = np.zeros([nStuTakExam])
a = np.zeros([nStuTakExam])
clsName = np.array(a, dtype=np.unicode)
for i in range(nStuTakExam):
    score[i] = tableXLS.cell(i + startRowIdx, scoreColumnIdx).value
    clsName[i] = tableXLS.cell(i + startRowIdx, classNameColumnIdx).value
#统计有多少个小班
b = np.zeros([nMaxClassNum])
cls = np.array(b, dtype=np.unicode)
cls[0] = clsName[0]
clsNum = 1
for i in range(1,nStuTakExam):
    for j in range(clsNum):
        found = 0
        if clsName[i] == cls[j]:
            found = 1
    if (found==0):
        clsNum+=1
        cls[clsNum-1] = clsName[i]
#分班统计人数
nMaxStuNum = 40
nStuNumPerClass = np.array(np.zeros([clsNum], dtype=np.int16))
nStuCountBelow60PerClass = np.array(np.zeros([clsNum], dtype=np.int16))
nStuCountAE90PerClass = np.array(np.zeros([clsNum], dtype=np.int16))
for i in range(nStuTakExam):
    for j in range(clsNum):
        if( clsName[i] == cls[j] ): nStuNumPerClass[j] += 1

#分小班统计各班的平均分，标准差，90分以上人数，不及格人数
maxStuNumPerClass = int(np.max(nStuNumPerClass))
scorePerClass = [[0.0]*maxStuNumPerClass for i in range(clsNum)]
scorePower2PerClass = [[0.0]*maxStuNumPerClass for i in range(clsNum)] #为计算标准差
c = np.zeros([clsNum])
idxStuNumPerClass = np.array(c, dtype=np.int16)
nStuCountBelow60 = 0
nStuCountAE90 = 0
for i in range(nStuTakExam):
    if (score[i] < 60): nStuCountBelow60 += 1
    if (score[i] >= 90): nStuCountAE90 += 1
    for j in range(clsNum):
        if (clsName[i] == cls[j]):
            scorePerClass[j][idxStuNumPerClass[j]] = score[i]
            scorePower2PerClass[j][idxStuNumPerClass[j]] = score[i]*score[i]
            if( score[i]<60 ): nStuCountBelow60PerClass[j]+=1
            if (score[i]>=90): nStuCountAE90PerClass[j] += 1
            idxStuNumPerClass[j] += 1

scoreAvgPerClass = np.zeros([clsNum])
scoreSTDPerClass = np.zeros([clsNum])
for j in range(clsNum):
    scoreAvgPerClass[j] = np.sum(scorePerClass[j])/nStuNumPerClass[j]
for j in range(clsNum):
    for k in range(nStuNumPerClass[j]):
        scoreSTDPerClass[j] = np.sqrt(np.sum(scorePower2PerClass[j])/nStuNumPerClass[j] - scoreAvgPerClass[j]*scoreAvgPerClass[j])

print "前十个学生成绩，供检查:", score[0:10]
scoreAvg = round(np.mean(score),2)
scoreDiffcult = round(scoreAvg/100.0,2)
scoreStD = round(np.std(score),2)
scoreMax = np.max(score)
scoreMin = np.min(score)
scorePhase = [0,30,35,40,45,50,55,60,65,70,75,80,85,90,95,100.01]
nScorePhase = len(scorePhase)-1
nStuCountPerScorePhase = np.zeros([nScorePhase])
for i in range(nStuTakExam):
    for j in range(nScorePhase):
        if score[i]>= scorePhase[j] and score[i]<scorePhase[j+1]:
            nStuCountPerScorePhase[j]+=1
print scoreAvg,scoreDiffcult,scoreStD
print scorePhase
print nStuCountPerScorePhase

fig, ax = plt.subplots()
fs = 28
gp.drawColumnChat( ax, scorePhasePlot, nStuCountPerScorePhase, xtickLable, nStuCountPerScorePhase, \
                   wordGraph_Title, wordGraph_XLable, wordGraph_YLable, fs, 3, bIntDisp=True)

document = Document(wordTemplateFileName)
tableWord = document.tables[0]
nrows = len(tableWord.rows)
rowFinish = 0
for row in range(nrows):
    ncells = len(tableWord.rows[row].cells)
    for k in range(ncells):
        if (tableWord.rows[row].cells[k].text == u"学生总数" and tableWord.rows[row].cells[k+1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % stuNum + u"人"
            break
        else:
            k+=1
    for k in range(k+2,ncells):
        if (tableWord.rows[row].cells[k].text == u"参考人数" and tableWord.rows[row].cells[k+1].text == u""):
          #  print len(tableWord.rows[row].cells[k+1].text)
            tableWord.rows[row].cells[k+1].text = '%r' % nStuTakExam + u"人"
            rowFinish = 1
            break
        else:
            k+=1

    if rowFinish == 1:
        break
rowFinish = 0
nextRow = row + 1
for row in range(nextRow, nrows):
    ncells = len(tableWord.rows[row].cells)
    for k in range(k+2, ncells):
       # print len(tableWord.rows[row].cells[k + 1].text)
        if (tableWord.rows[row].cells[k].text == u"缓考人数" and tableWord.rows[row].cells[k+1].text==u""):
            tableWord.rows[row].cells[k+1].text = '%r' % nDelayExamStu + u"人"
            break
        else:
            k+=1
    for k in range(ncells):
        if (tableWord.rows[row].cells[k].text == u"旷考人数" and tableWord.rows[row].cells[k+1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % nNotAttendExamStu + u"人"
            break
        else:
            k+=1
    for k in range(k+2, ncells):
        if (tableWord.rows[row].cells[k].text == u"违纪人数" and tableWord.rows[row].cells[k+1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % nViolationExamStu + u"人"
            rowFinish = 1
            break
        else:
            k+=1
    if rowFinish==1:
        break
#成绩分布表格
nextRow = row + 4
for row in range(nextRow, nrows):
    i = 0
    foundScorePhase = 0
    ncells = len(tableWord.rows[row].cells)
    for k in range(ncells):
        if (tableWord.rows[row].cells[k].text == scorePhaseString[i]):
            foundScorePhase = 1
            tableWord.rows[row+1].cells[k].text = '%d' % nStuCountPerScorePhase[i]
            i+=1
    if(foundScorePhase==1):
        for k in range(0,len(tableWord.rows[row+2].cells)-1):
            if (i<len(scorePhaseString) and tableWord.rows[row+2].cells[k].text == scorePhaseString[i]):
                tableWord.rows[row+3].cells[k].text = '%d' % nStuCountPerScorePhase[i]
                i += 1
        break
rowFinish=0
nextRow = row + 4
for row in range(nextRow, nrows):
    ncells = len(tableWord.rows[row].cells)
    for k in range(ncells):
        if (tableWord.rows[row].cells[k].text == u"平均分" and tableWord.rows[row].cells[k+1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % scoreAvg
            break
        else:
            k+=1
    for k in range(k+2,ncells):
        if (tableWord.rows[row].cells[k].text == u"标准差" and tableWord.rows[row].cells[k + 1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % scoreStD
            break
        else:
            k+=1
    for k in range(k+2,ncells):
        if (tableWord.rows[row].cells[k].text == u"最高分" and tableWord.rows[row].cells[k + 1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % scoreMax
            break
        else:
            k+=1
    for k in range(k+2,ncells):
        if (tableWord.rows[row].cells[k].text == u"最低分" and tableWord.rows[row].cells[k + 1].text == u""):
            tableWord.rows[row].cells[k+1].text = '%r' % scoreMin
            break
        else:
            k+=1
    if rowFinish==1:
        break;
#小班分析
rowFinish=0
row = 11
ncells = len(tableWord.rows[row].cells)
nextK = 0
for n in range(0, clsNum):
    for k in range(nextK, ncells):
        if (tableWord.rows[row].cells[k].text == str(n).decode('utf8')):# and tableWord.rows[row].cells[k+1].text == u""):
            tableWord.rows[row+1].cells[k].text = '%r' % round(scoreAvgPerClass[n],2)
            tableWord.rows[row + 2].cells[k].text = '%r' % round(scoreSTDPerClass[n],2)
            tableWord.rows[row + 3].cells[k].text = '%r' % nStuCountAE90PerClass[n]
            tableWord.rows[row + 4].cells[k].text = '%r' % nStuCountBelow60PerClass[n]
            nextK = k+1
            break
        else:
            k+=1
#page 2
tableWordPage2 = document.tables[1]
nrowsPage2 = len(tableWordPage2.rows)
rowFinish = 0
for row in range(nrowsPage2):
    if (u"试卷整体合理性" in tableWordPage2.rows[row].cells[0].text):
        tableWordPage2.rows[row].cells[0].text += u"\n"
        tableWordPage2.rows[row].cells[0].text += u"2. 试卷难度：难度系数为"+str(scoreDiffcult)
        tableWordPage2.rows[row].cells[0].text += u"，合班平均分为"
        tableWordPage2.rows[row].cells[0].text += str(scoreAvg)
        tableWordPage2.rows[row].cells[0].text += u"，"
        tableWordPage2.rows[row].cells[0].text += u"试卷难度"
        strDiffGrad = u""
        if(scoreDiffcult<0.7):
            strDiffGrad = u"较难"
        elif(scoreDiffcult<0.85):
            strDiffGrad = u"适中"
        else:
            strDiffGrad = u"较易"
        tableWordPage2.rows[row].cells[0].text += strDiffGrad
        tableWordPage2.rows[row].cells[0].text += u"。\n"

        #if (u"不及格学生的试卷分析" in tableWordPage2.rows[row].cells[0].text):
        tableWordPage2.rows[row].cells[0].text += u"4. 成绩直方图分析:\n"
        tableWordPage2.rows[row].cells[0].text += u"成绩直方图趋于正态分布，60分以下" + str(nStuCountBelow60)
        tableWordPage2.rows[row].cells[0].text += u"人，约占本合班参考人数的" + str(round(nStuCountBelow60*100.0/nStuTakExam,2))
        tableWordPage2.rows[row].cells[0].text += u"%，"
        tableWordPage2.rows[row].cells[0].text += u"90分以上" + str(nStuCountAE90)
        tableWordPage2.rows[row].cells[0].text += u"人，约占本合班参考人数的" + str(round(nStuCountAE90*100.0/nStuTakExam,2))
        tableWordPage2.rows[row].cells[0].text += u"%。"
        break
document.save('.\\data\\scoreAnalysisResult.docx')
print "请手动保存显示的成绩分析图！！！"
plt.show()




